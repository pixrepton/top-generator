#!/usr/bin/env python3
"""
PDF to DOCX Converter
Szybki lokalny konwerter plików PDF na DOCX
"""

import os
import sys
import argparse
import time
from pathlib import Path
from concurrent.futures import ProcessPoolExecutor, as_completed
from tqdm import tqdm
import multiprocessing

# Import bibliotek do konwersji
import pdfplumber
from docx import Document
from docx.shared import Inches

class PDFToDocxConverter:
    def __init__(self, max_workers=None):
        """
        Inicjalizuje konwerter
        
        Args:
            max_workers: Maksymalna liczba procesów równoległych (domyślnie: liczba CPU)
        """
        self.max_workers = max_workers or min(multiprocessing.cpu_count(), 8)
        
    def convert_batch(self, input_paths, output_dir=None):
        """Konwertuje wiele plików PDF na DOCX"""
        results = {}
        for pdf_path in input_paths:
            success, output_path, error = self.convert_single_file(pdf_path, output_dir)
            results[pdf_path] = (success, output_path, error)
        return results

    def convert_single_file(self, pdf_path, output_dir=None, custom_name=None):
        """
        Konwertuje pojedynczy plik PDF na DOCX
        
        Args:
            pdf_path: Ścieżka do pliku PDF
            output_dir: Katalog wyjściowy (domyślnie: katalog źródłowy)
            custom_name: Niestandardowa nazwa pliku wyjściowego
            
        Returns:
            tuple: (success: bool, output_path: str, error_message: str)
        """
        try:
            pdf_path = Path(pdf_path)
            
            if not pdf_path.exists():
                return False, "", f"Plik nie istnieje: {pdf_path}"
            
            if not pdf_path.suffix.lower() == '.pdf':
                return False, "", f"Plik nie jest PDF: {pdf_path}"
            
            # Określ katalog wyjściowy
            if output_dir:
                output_dir = Path(output_dir)
                output_dir.mkdir(parents=True, exist_ok=True)
            else:
                output_dir = pdf_path.parent
            
            # Określ nazwę pliku wyjściowego
            if custom_name:
                output_name = custom_name
                if not output_name.endswith('.docx'):
                    output_name += '.docx'
            else:
                output_name = pdf_path.stem + '.docx'
            
            output_path = output_dir / output_name
            
            # Konwersja PDF do DOCX
            doc = Document()
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    if page_num > 0:
                        doc.add_page_break()
                    
                    # Dodaj nagłówek strony
                    doc.add_heading(f'Strona {page_num + 1}', level=2)
                    
                    # Wyciągnij tekst ze strony
                    text = page.extract_text()
                    
                    if text:
                        # Podziel tekst na akapity
                        paragraphs = text.split('\n\n')
                        for paragraph in paragraphs:
                            if paragraph.strip():
                                doc.add_paragraph(paragraph.strip())
                    else:
                        doc.add_paragraph("[Nie udało się wyciągnąć tekstu z tej strony]")
                    
                    # Wyciągnij tabele jeśli istnieją
                    tables = page.extract_tables()
                    if tables:
                        for table_data in tables:
                            if table_data and len(table_data) > 0:
                                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                                table.style = 'Table Grid'
                                
                                for row_idx, row_data in enumerate(table_data):
                                    for col_idx, cell_data in enumerate(row_data):
                                        if cell_data:
                                            table.cell(row_idx, col_idx).text = str(cell_data)
            
            # Zapisz dokument
            doc.save(output_path)
            
            return True, str(output_path), ""
            
        except Exception as e:
            return False, "", f"Błąd konwersji: {str(e)}"

def main():
    if len(sys.argv) < 2:
        print("Użycie: python pdf_to_docx_converter.py <plik.pdf>")
        sys.exit(1)
    
    pdf_file = sys.argv[1]
    converter = PDFToDocxConverter()
    
    print("🔄 PDF TO DOCX CONVERTER")
    print("="*40)
    
    success, output_path, error = converter.convert_single_file(pdf_file)
    
    if success:
        print(f"✅ Konwersja zakończona: {output_path}")
    else:
        print(f"❌ Błąd konwersji: {error}")

if __name__ == "__main__":
    main()

def find_pdf_files(directory):
    """Znajduje wszystkie pliki PDF w katalogu i podkatalogach"""
    pdf_files = []
    directory = Path(directory)
    
    if directory.is_file() and directory.suffix.lower() == '.pdf':
        return [str(directory)]
    
    for pdf_file in directory.rglob('*.pdf'):
        pdf_files.append(str(pdf_file))
    
    return pdf_files

def print_results(results):
    """Wyświetla podsumowanie wyników konwersji"""
    print("\n" + "="*60)
    print("📊 PODSUMOWANIE KONWERSJI")
    print("="*60)
    
    successful = sum(1 for success, _, _ in results.values() if success)
    failed = len(results) - successful
    
    print(f"✅ Pomyślne konwersje: {successful}")
    print(f"❌ Nieudane konwersje: {failed}")
    print(f"📁 Łącznie plików: {len(results)}")
    
    if failed > 0:
        print("\n❌ Błędy konwersji:")
        for path, (success, output_path, error) in results.items():
            if not success:
                print(f"   • {Path(path).name}: {error}")
    
    if successful > 0:
        print(f"\n✅ Pomyślnie skonwertowane pliki:")
        for path, (success, output_path, error) in results.items():
            if success:
                print(f"   • {Path(path).name} → {Path(output_path).name}")
