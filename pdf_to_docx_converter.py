#!/usr/bin/env python3
"""
PDF to DOCX Converter
Szybki lokalny konwerter plików PDF na DOCX
"""

import os
import sys
import argparse
import time
from pathlib import Path
from concurrent.futures import ProcessPoolExecutor, as_completed
from tqdm import tqdm
import multiprocessing

# Import bibliotek do konwersji
import pdfplumber
from docx import Document
from docx.shared import Inches

class PDFToDocxConverter:
    def __init__(self, max_workers=None):
        """
        Inicjalizuje konwerter
        
        Args:
            max_workers: Maksymalna liczba procesów równoległych (domyślnie: liczba CPU)
        """
        self.max_workers = max_workers or min(multiprocessing.cpu_count(), 8)
        
    def convert_single_file(self, pdf_path, output_dir=None, custom_name=None):
        """
        Konwertuje pojedynczy plik PDF na DOCX
        
        Args:
            pdf_path: Ścieżka do pliku PDF
            output_dir: Katalog wyjściowy (domyślnie: katalog źródłowy)
            custom_name: Niestandardowa nazwa pliku wyjściowego
            
        Returns:
            tuple: (success: bool, output_path: str, error_message: str)
        """
        try:
            pdf_path = Path(pdf_path)
            
            if not pdf_path.exists():
                return False, "", f"Plik nie istnieje: {pdf_path}"
            
            if not pdf_path.suffix.lower() == '.pdf':
                return False, "", f"Plik nie jest PDF: {pdf_path}"
            
            # Określ katalog wyjściowy
            if output_dir:
                output_dir = Path(output_dir)
                output_dir.mkdir(parents=True, exist_ok=True)
            else:
                output_dir = pdf_path.parent
            
            # Określ nazwę pliku wyjściowego
            if custom_name:
                output_name = custom_name
                if not output_name.endswith('.docx'):
                    output_name += '.docx'
            else:
                output_name = pdf_path.stem + '.docx'
            
            output_path = output_dir / output_name
            
            # Konwersja PDF do DOCX
            doc = Document()
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    if page_num > 0:
                        doc.add_page_break()
                    
                    # Dodaj nagłówek strony
                    doc.add_heading(f'Strona {page_num + 1}', level=2)
                    
                    # Wyciągnij tekst ze strony
                    text = page.extract_text()
                    
                    if text:
                        # Podziel tekst na akapity
                        paragraphs = text.split('\n\n')
                        for paragraph in paragraphs:
                            if paragraph.strip():
                                doc.add_paragraph(paragraph.strip())
                    else:
                        doc.add_paragraph("[Nie udało się wyciągnąć tekstu z tej strony]")
                    
                    # Wyciągnij tabele jeśli istnieją
                    tables = page.extract_tables()
                    if tables:
                        for table_data in tables:
                            if table_data and len(table_data) > 0:
                                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                                table.style = 'Table Grid'
                                
                                for row_idx, row_data in enumerate(table_data):
                                    for col_idx, cell_data in enumerate(row_data):
                                        if cell_data:
                                            table.cell(row_idx, col_idx).text = str(cell_data)
            
            # Zapisz dokument
            doc.save(output_path)
            
            return True, str(output_path), ""
            
        except Exception as e:
            return False, "", f"Błąd konwersji: {str(e)}"
    
    def convert_batch(self, input_paths, output_dir=None, progress_callback=None):
        """
        Konwertuje wiele plików PDF na DOCX równolegle
        
        Args:
            input_paths: Lista ścieżek do plików PDF
            output_dir: Katalog wyjściowy
            progress_callback: Funkcja callback dla postępu
            
        Returns:
            dict: Wyniki konwersji {path: (success, output_path, error)}
        """
        results = {}
        
        with ProcessPoolExecutor(max_workers=self.max_workers) as executor:
            # Przygotuj zadania
            future_to_path = {}
            for pdf_path in input_paths:
                future = executor.submit(self.convert_single_file, pdf_path, output_dir)
                future_to_path[future] = pdf_path
            
            # Przetwarzaj wyniki
            with tqdm(total=len(input_paths), desc="Konwertowanie plików") as pbar:
                for future in as_completed(future_to_path):
                    pdf_path = future_to_path[future]
                    try:
                        success, output_path, error = future.result()
                        results[pdf_path] = (success, output_path, error)
                    except Exception as e:
                        results[pdf_path] = (False, "", f"Błąd wykonania: {str(e)}")
                    
                    pbar.update(1)
                    if progress_callback:
                        progress_callback(pdf_path, results[pdf_path])
        
        return results

def find_pdf_files(directory):
    """Znajduje wszystkie pliki PDF w katalogu i podkatalogach"""
    pdf_files = []
    directory = Path(directory)
    
    if directory.is_file() and directory.suffix.lower() == '.pdf':
        return [str(directory)]
    
    for pdf_file in directory.rglob('*.pdf'):
        pdf_files.append(str(pdf_file))
    
    return pdf_files

def print_results(results):
    """Wyświetla podsumowanie wyników konwersji"""
    print("\n" + "="*60)
    print("📊 PODSUMOWANIE KONWERSJI")
    print("="*60)
    
    successful = sum(1 for success, _, _ in results.values() if success)
    failed = len(results) - successful
    
    print(f"✅ Pomyślne konwersje: {successful}")
    print(f"❌ Nieudane konwersje: {failed}")
    print(f"📁 Łącznie plików: {len(results)}")
    
    if failed > 0:
        print("\n❌ Błędy konwersji:")
        for path, (success, output_path, error) in results.items():
            if not success:
                print(f"   • {Path(path).name}: {error}")
    
    if successful > 0:
        print(f"\n✅ Pomyślnie skonwertowane pliki:")
        for path, (success, output_path, error) in results.items():
            if success:
                print(f"   • {Path(path).name} → {Path(output_path).name}")

def main():
    parser = argparse.ArgumentParser(
        description="Konwerter PDF do DOCX - szybka konwersja lokalnie",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Przykłady użycia:
  %(prog)s dokument.pdf                    # Konwertuj pojedynczy plik
  %(prog)s /ścieżka/do/katalogu/           # Konwertuj wszystkie PDF w katalogu
  %(prog)s plik.pdf -o /wyjście/           # Określ katalog wyjściowy
  %(prog)s katalog/ -w 4                   # Użyj 4 procesów równoległych
  %(prog)s plik.pdf -n "nowa_nazwa"        # Określ niestandardową nazwę
        """
    )
    
    parser.add_argument('input', help='Plik PDF lub katalog z plikami PDF')
    parser.add_argument('-o', '--output', help='Katalog wyjściowy (domyślnie: katalog źródłowy)')
    parser.add_argument('-n', '--name', help='Niestandardowa nazwa dla pojedynczego pliku')
    parser.add_argument('-w', '--workers', type=int, help='Liczba procesów równoległych')
    parser.add_argument('-v', '--verbose', action='store_true', help='Szczegółowe informacje')
    
    args = parser.parse_args()
    
    print("🔄 PDF TO DOCX CONVERTER")
    print("="*40)
    
    # Sprawdź czy wejście istnieje
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"❌ Ścieżka nie istnieje: {args.input}")
        sys.exit(1)
    
    # Znajdź pliki PDF
    if input_path.is_file():
        if input_path.suffix.lower() != '.pdf':
            print(f"❌ Plik nie jest PDF: {args.input}")
            sys.exit(1)
        pdf_files = [str(input_path)]
    else:
        pdf_files = find_pdf_files(input_path)
        if not pdf_files:
            print(f"❌ Nie znaleziono plików PDF w: {args.input}")
            sys.exit(1)
    
    print(f"📁 Znaleziono {len(pdf_files)} plików PDF")
    
    # Inicjalizuj konwerter
    converter = PDFToDocxConverter(max_workers=args.workers)
    
    start_time = time.time()
    
    if len(pdf_files) == 1 and args.name:
        # Konwersja pojedynczego pliku z niestandardową nazwą
        success, output_path, error = converter.convert_single_file(
            pdf_files[0], args.output, args.name
        )
        
        if success:
            print(f"✅ Konwersja zakończona: {output_path}")
        else:
            print(f"❌ Błąd konwersji: {error}")
            sys.exit(1)
    else:
        # Konwersja wsadowa
        def progress_callback(path, result):
            if args.verbose:
                success, output_path, error = result
                if success:
                    print(f"✅ {Path(path).name} → {Path(output_path).name}")
                else:
                    print(f"❌ {Path(path).name}: {error}")
        
        results = converter.convert_batch(
            pdf_files, 
            args.output,
            progress_callback if args.verbose else None
        )
        
        print_results(results)
    
    end_time = time.time()
    print(f"\n⏱️  Czas wykonania: {end_time - start_time:.2f} sekund")

if __name__ == "__main__":
    main()